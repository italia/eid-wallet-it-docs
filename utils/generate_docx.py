#!/usr/bin/env python3
import os
import sys
import csv
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Missing dependency: python-docx. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(2)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    para = document.add_paragraph()
    run = para.add_run(text)
    if level == 1:
        run.bold = True
        run.font.size = Pt(16)
    elif level == 2:
        run.bold = True
        run.font.size = Pt(14)
    else:
        run.bold = True
        run.font.size = Pt(12)


def add_body_text(document: Document, text: str) -> None:
    for line in text.splitlines():
        document.add_paragraph(line)


def add_tsv_table(document: Document, tsv_path: str) -> None:
    with open(tsv_path, "r", encoding="utf-8") as f:
        reader = csv.reader(f, delimiter='\t')
        rows = list(reader)

    if not rows:
        return

    header = rows[0]
    table = document.add_table(rows=1, cols=len(header))
    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(header):
        hdr_cells[idx].text = col

    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, col in enumerate(row):
            cells[idx].text = col


def main():
    repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    rst_path = os.path.join(repo_root, "docs", "it", "attestati-interesse-pubblico.rst")
    tsv_path = os.path.join(repo_root, "attestazioni_test_matrix.tsv")
    out_path = os.path.join(repo_root, "Attestati-Interesse-Pubblico.docx")

    if not os.path.exists(rst_path):
        print(f"RST not found: {rst_path}", file=sys.stderr)
        sys.exit(1)
    if not os.path.exists(tsv_path):
        print(f"TSV not found: {tsv_path}", file=sys.stderr)
        sys.exit(1)

    with open(rst_path, "r", encoding="utf-8") as f:
        rst_content = f.read()

    doc = Document()
    title = "Attestati di interesse pubblico e di attributi pubblici"
    add_heading(doc, title, level=1)
    subtitle = f"Generato automaticamente il {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    p = doc.add_paragraph(subtitle)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_heading(doc, "Contenuti (estratto dalla versione RST)", level=2)
    add_body_text(doc, rst_content)

    add_heading(doc, "Matrice di test (TSV)", level=2)
    add_tsv_table(doc, tsv_path)

    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()









